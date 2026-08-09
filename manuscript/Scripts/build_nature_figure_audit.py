"""Create synchronized Nature-style figure/evidence audit DOCX and Markdown."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "nature_audit"
STEM = "Nature_ReviewStudio_figure_evidence_audit_v65_20260809"


SECTIONS = [
    (
        "审稿人1 — 图表与信息设计",
        "整体评价",
        "修订后的图表层级已明显改善。原Mn-oxide两栏Figure 6没有独立视觉价值，现已由紧凑的Table 4取代；held-out图已移除数据区内的图例和结果文字，逐seed差值与停止预算结果可以直接读取。当前主要剩余图像风险是Figure 1原始位图中最小文字的最终印刷字号。",
        [
            ("Mn-oxide控制的表达形式", "已完成。主文Section 7.1删除原Figure 6，Table 4集中报告10/10配对AUTC差值为0、批次候选集合完全一致、199次direct、1次correction和0次有效替换。该表比重合曲线更直接地支撑fallback结论。"),
            ("Held-out图的数据遮挡", "已完成。重新编号后的Figure 7a使用以零为基准的signed stem plot；图中不再放置‘6/10’、均值结果文字或类别图例。Figure 7b的共享图例位于坐标区上方，未覆盖误差条或数据点。"),
        ],
        [
            ("Figure 1最小文字", "4096×2730、600 dpi投稿导出已经生成，但上采样不能替代原生矢量文字。若要严格证明所有文字达到8 pt，仍需矢量源文件或拆分流程图。")
        ],
        "高",
    ),
    (
        "审稿人2 — 统计与方法证据",
        "整体评价",
        "新的seeds 15–24冻结held-out比较解决了旧档案查询前缀重复、无法作为独立重复的问题。主文同时报告平均提升、bootstrap区间、exact Wilcoxon结果和逐seed方向，避免把平均优势误写成每个seed均占优。",
        [
            ("独立轨迹与统计推断", "已完成。每种方法均有10个不同的完整序列哈希。Gate的平均AUTC160差值为+0.0256（相对Greedy为8.6%），bootstrap区间为[0.0067, 0.0449]；exact two-sided Wilcoxon p=0.105，6/10个seed支持Gate。因此正文只支持平均早期优势，不声称普遍seed级优势。"),
            ("Full Gate与Group-only归因", "已完成边界修正。正文明确两者保存序列相同，Margin-only与Greedy一致；方法优势归因于group-concentration触发的correction route，不再声称margin条件具有独立效果。"),
        ],
        [
            ("DFT-evaluability样本量", "该模型仅基于20次历史DFT尝试，仍只能作为acquisition-blind、post-selection、exploratory指标；所有相关数量必须继续标注为ML-estimated或predicted。")
        ],
        "高",
    ),
    (
        "审稿人3 — 可复现性与主张边界",
        "整体评价",
        "修订稿已经补齐主文、SI、图表源数据、重绘脚本、结构文件、参考文献核验和SHA-256清单。论文对离线ALIGNN回放、DFT评估和热力学稳定性的边界总体清楚，但投稿时不得将其重新包装为prospective DFT active learning或稳定新材料发现。",
        [
            ("离线代理任务与真实DFT", "已完成主张边界修正。ALIGNN代理标签、MLIP能量和PAW–PBE+U形成能保持为不同能量约定；现有四个LiCr2O4结果仅证明在给定协议下获得了收敛且互不重复的松弛结构。"),
            ("源文件与追溯性", "已完成。定量图均可追溯到CSV及Python重绘脚本，SI表格保留逐seed结果、路由记录和模型审计；投稿ZIP可在空目录重新编译。"),
        ],
        [
            ("未完成的前瞻验证", "未执行prospective DFT acquisition loop、energy-above-hull和广泛磁态搜索。这些属于明确保留的研究边界，而不是已解决实验。")
        ],
        "高",
    ),
]


TASKS = [
    ("T1", "图表", "以Table 4替换旧Figure 6", "DONE", "否"),
    ("T2", "图表", "重绘held-out图并移除数据区遮挡", "DONE", "否"),
    ("T3", "统计", "保留平均优势与Wilcoxon不显著之间的平衡表述", "DONE", "否"),
    ("T4", "证据边界", "将n=20 evaluability结果保持为探索性估计", "DONE", "否"),
    ("T5", "图像规范", "获取Figure 1矢量源或拆分以严格满足8 pt", "TODO_AUTHOR_CONFIRM", "否"),
]


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("Nature-style figure and evidence audit"), 18, True, "1F4D78")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Energy-Gated DA-TPP修订稿 · 2026-08-09"), 11, False, "555555")

    p = doc.add_paragraph()
    set_font(p.add_run("修订等级：Accept with minor revisions。"), 11, True)
    set_font(p.add_run(" 旧Figure 6和Figure 8的主要信息设计问题已经关闭；剩余事项主要是Figure 1原始位图的最小字号，以及必须继续保持的科学证据边界。"), 11)

    for reviewer, overall_label, overall, majors, minors, confidence in SECTIONS:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        set_font(h.add_run(reviewer), 14, True, "2E74B5")
        p = doc.add_paragraph()
        set_font(p.add_run(f"{overall_label}。"), 10.5, True)
        set_font(p.add_run(overall), 10.5)
        p = doc.add_paragraph()
        set_font(p.add_run("主要意见"), 11, True, "1F4D78")
        for idx, (heading, body) in enumerate(majors, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            set_font(p.add_run(f"{idx}. {heading} — "), 10.5, True)
            set_font(p.add_run(body), 10.5)
        p = doc.add_paragraph()
        set_font(p.add_run("次要意见"), 11, True, "1F4D78")
        for idx, (heading, body) in enumerate(minors, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            set_font(p.add_run(f"{idx}. {heading} — "), 10.5, True)
            set_font(p.add_run(body), 10.5)
        p = doc.add_paragraph()
        set_font(p.add_run(f"置信度：{confidence}"), 10.5, True)

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    set_font(h.add_run("跨审稿共识"), 14, True, "2E74B5")
    p = doc.add_paragraph()
    set_font(p.add_run("当前版本可以作为透明、边界明确的CMC投稿稿件。"), 10.5, True)
    set_font(p.add_run(" 图表已不再重复呈现同一数据，核心Gate–Greedy推断使用独立held-out轨迹；但稿件不能声称完成了prospective DFT active learning、热力学相稳定性验证或基于20条历史记录的确认性DFT-evaluability结论。"), 10.5)

    h = doc.add_paragraph()
    set_font(h.add_run("修订任务表"), 14, True, "2E74B5")
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    widths = [0.55, 0.85, 3.45, 1.25, 0.4]
    headers = ["ID", "类别", "任务", "状态", "阻断"]
    for idx, (cell, width, text) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "F2F4F7")
        cell.text = ""
        set_font(cell.paragraphs[0].add_run(text), 9.5, True)
    for row in TASKS:
        cells = table.add_row().cells
        for cell, width, text in zip(cells, widths, row):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            set_font(cell.paragraphs[0].add_run(text), 9.5)
    doc.save(path)


def build_markdown(path: Path):
    lines = [
        "# Nature-style figure and evidence audit",
        "",
        "**修订等级：Accept with minor revisions。**旧Figure 6和Figure 8的主要信息设计问题已经关闭；剩余事项主要是Figure 1原始位图的最小字号，以及必须继续保持的科学证据边界。",
        "",
    ]
    for reviewer, overall_label, overall, majors, minors, confidence in SECTIONS:
        lines += [f"## {reviewer}", "", f"**{overall_label}。** {overall}", "", "### 主要意见", ""]
        for idx, (heading, body) in enumerate(majors, 1):
            lines.append(f"{idx}. **{heading}** — {body}")
        lines += ["", "### 次要意见", ""]
        for idx, (heading, body) in enumerate(minors, 1):
            lines.append(f"{idx}. **{heading}** — {body}")
        lines += ["", f"**置信度：{confidence}**", ""]
    lines += [
        "## 跨审稿共识",
        "",
        "当前版本可以作为透明、边界明确的CMC投稿稿件。图表已不再重复呈现同一数据，核心Gate–Greedy推断使用独立held-out轨迹；但稿件不能声称完成了prospective DFT active learning、热力学相稳定性验证或基于20条历史记录的确认性DFT-evaluability结论。",
        "",
        "## 修订任务表",
        "",
        "| ID | 类别 | 任务 | 状态 | 阻断 |",
        "|---|---|---|---|---|",
    ]
    lines += ["| " + " | ".join(row) + " |" for row in TASKS]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_docx(OUT / f"{STEM}.docx")
    build_markdown(OUT / f"{STEM}.md")


if __name__ == "__main__":
    main()
